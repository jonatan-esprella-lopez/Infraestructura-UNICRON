from docx import Document
import sys, os

path = r'C:\Users\rolda\source\Projects\cocha-tech\Infraestructura-UNICRON\docs\legalities'
files = sys.argv[1:]

for f in files:
    filepath = os.path.join(path, f)
    try:
        doc = Document(filepath)
        print(f'=== {f} ===')
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            print('\n[TABLE]')
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end='\t')
                print()
        print()
    except Exception as e:
        print(f'=== {f} ===\nError: {e}\n')
